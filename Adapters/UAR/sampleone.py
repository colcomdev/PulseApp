from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime

# ================= SETTINGS =================
excel_path = "sales.xlsx"
word_path = "Company_Report.docx"

company_name = "ABC Corporation"
company_address = "123 Business Avenue, New York, NY 10001"
company_phone = "Phone: (123) 456-7890"

# Form values (filled programmatically)
report_number = "RPT-2026-001"
prepared_by = "John Doe"
department = "Finance"
# ============================================


# Load Excel
wb = load_workbook(excel_path)
ws = wb.active

# Create Word document
doc = Document()
section = doc.sections[0]

# ==================================================
# HEADER WITH 3 SECTIONS (Left | Center | Right)
# ==================================================

header = section.header
header_para = header.paragraphs[0]

# Use tab stops to create 3 aligned sections
header_para.text = company_name
header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

header_para.add_run("\t" + "OFFICIAL REPORT").bold = True
header_para.add_run("\t" + datetime.today().strftime("%B %d, %Y"))

# ==================================================
# FOOTER
# ==================================================

footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = f"{company_address} | {company_phone}"
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ==================================================
# TITLE
# ==================================================

doc.add_heading("Employee Data Report", level=1)

# ==================================================
# FORM SECTION (Programmatically Filled)
# ==================================================

doc.add_paragraph("Report Number: " + report_number)
doc.add_paragraph("Prepared By: " + prepared_by)
doc.add_paragraph("Department: " + department)
doc.add_paragraph("Date Generated: " + datetime.today().strftime("%Y-%m-%d"))

doc.add_paragraph("\n")

# ==================================================
# INSERT EXCEL DATA INTO WORD TABLE
# ==================================================

rows = ws.max_row
cols = ws.max_column

table = doc.add_table(rows=rows, cols=cols)
table.style = "Table Grid"

for i, row in enumerate(ws.iter_rows(values_only=True)):
    for j, cell in enumerate(row):
        table.rows[i].cells[j].text = str(cell)

# Save document
doc.save(word_path)

print("Word document created successfully!")