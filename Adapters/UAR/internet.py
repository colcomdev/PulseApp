from docx import Document
from docx.shared import Inches,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document()

# Access header
section = doc.sections[0]
header = section.header

# Create table with 1 row and 3 columns
table = header.add_table(rows=1, cols=3, width=Inches(6))

# LEFT LOGO
left_cell = table.cell(0,0)
left_para = left_cell.paragraphs[0]
left_run = left_para.add_run()
left_run.add_picture("logo1.png", width=Inches(1))
left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

# CENTER TEXT
center_cell = table.cell(0,1)
center_para = center_cell.paragraphs[0]
center_para.text = "COLCOM FOODS LIMITED\nInformation Technology Department"
center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# RIGHT LOGO
right_cell = table.cell(0,2)
right_para = right_cell.paragraphs[0]
right_run = right_para.add_run()
right_run.add_picture("logo2.png", width=Inches(2))
right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Add body text
doc.add_paragraph("Document Content Starts Here...")
# Create main table
table = doc.add_table(rows=8, cols=4)
table.style = "Table Grid"

# ---- Row 1: Services Required Title ----
title_cell = table.cell(0,0).merge(table.cell(0,3))
title_cell.text = "Services Required"

# ---- Left column merged for Tick Services ----
left = table.cell(1,0).merge(table.cell(4,0))
left.text = "Tick Services Required"

# ---- Service options ----
table.cell(1,1).text = ""
table.cell(1,2).merge(table.cell(1,3)).text = "Internet"

table.cell(2,1).text = ""
table.cell(2,2).merge(table.cell(2,3)).text = "All facilities (Internet Mail, World Wide Web, News Group)"

# ---- ICT Office ONLY ----
ict = table.cell(3,1).merge(table.cell(3,3))
ict.text = "ICT Office ONLY"

# ---- VSAT Connectivity ----
table.cell(4,1).text = ""
table.cell(4,2).merge(table.cell(4,3)).text = "VSAT Connectivity (Reserved IPs)"

# ---- Device Header ----
table.cell(5,0).text = ""
table.cell(5,1).text = "Device"
table.cell(5,2).text = "IP Address"
table.cell(5,3).text = "Mac Address"

# ---- Empty rows for entries ----
for r in range(6,8):
    for c in range(4):
        table.cell(r,c).text = ""

# ---- Business Justification Section ----
doc.add_paragraph()

biz_table = doc.add_table(rows=2, cols=4)
biz_table.style = "Table Grid"

biz_table.cell(0,0).text = "Business Justification"
biz_table.cell(0,1).merge(biz_table.cell(0,3))

biz_table.cell(1,0).merge(biz_table.cell(1,3)).text = ""

#========================================================================================================================


# --- Applicant Declaration ---
p = doc.add_paragraph(
"I, the undersigned, acknowledge that I have read, understand and agree to comply "
"with the provisions of the Internet Policies of Colcom Foods Limited."
)

doc.add_paragraph().add_run("Signed by Applicant:").bold = True

# Applicant signature table
table1 = doc.add_table(rows=2, cols=3)
table1.style = "Table Grid"

table1.cell(0,0).text = "Signature"
table1.cell(0,1).text = "Name"
table1.cell(0,2).text = "Date"

# Empty row
for c in range(3):
    table1.cell(1,c).text = ""


# --- HOD Declaration ---
doc.add_paragraph()
doc.add_paragraph(
"I, as Head of Department, am satisfied that there is a business justification for "
"the above Internet service and that the applicant has read and agreed to comply "
"with the Internet Policies of Colcom Foods Limited."
)

doc.add_paragraph().add_run("Authorized by Head of Department").bold = True

# HOD signature table
table2 = doc.add_table(rows=2, cols=3)
table2.style = "Table Grid"

table2.cell(0,0).text = "Signature"
table2.cell(0,1).text = "Name"
table2.cell(0,2).text = "Date"

for c in range(3):
    table2.cell(1,c).text = ""


# --- IT Department Section ---
doc.add_paragraph()

it_table = doc.add_table(rows=4, cols=3)
it_table.style = "Table Grid"

# Header row merged
header = it_table.cell(0,0).merge(it_table.cell(0,2))
header.text = "For use by Information Technology Department"

# Apply black shading to header
tc = header._element
tcPr = tc.get_or_add_tcPr()
shading = OxmlElement('w:shd')
shading.set(qn('w:fill'), "000000")
tcPr.append(shading)

# White text
run = header.paragraphs[0].runs[0]
run.font.color.rgb = None


# Rows
it_table.cell(1,0).text = "Request Approved By (Name)"
it_table.cell(1,1).text = "Signature"
it_table.cell(1,2).text = "Date"

it_table.cell(2,0).text = "Work Completed by (Name)"
it_table.cell(2,1).text = "Signature"
it_table.cell(2,2).text = "Date"

it_table.cell(3,0).text = "Work Reviewed by (Name)"
it_table.cell(3,1).text = "Signature"
it_table.cell(3,2).text = "Date"





# ---------------- FOOTER ----------------
footer = section.footer
footer_table = footer.add_table(rows=1, cols=3, width=Inches(6))

# Left footer text
left_footer = footer_table.cell(0,0).paragraphs[0]
left_footer.text = "Colcom Foods Limited"
left_footer.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Center footer text
center_footer = footer_table.cell(0,1).paragraphs[0]
center_footer.text = "Internet User Access Request Form"
center_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Right footer page number
right_footer = footer_table.cell(0,2).paragraphs[0]
right_footer.text = "Page "
right_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.save("header_layout.docx")