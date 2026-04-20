from docx import Document

document = Document()
section = document.sections[0]
header = section.header
#paragraph = header.paragraphs[0]
#paragraph.text = "Title of my document"
paragraph = header.paragraphs[0]
paragraph.text = "Left Text\tCenter Text\tRight Text"
paragraph.style = document.styles["Header"]

table = document.add_table(3,3)
table.cell[0]="Test"




"""
=========================================================================================================================================================
INHERITABLE FOOTER DESIGNS
=========================================================================================================================================================
"""
document.save('test.docx')